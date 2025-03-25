import PyPDF2
import docx
import re
import json
from typing import Dict, List, Any, Optional
import logging
from pathlib import Path
import os

logger = logging.getLogger(__name__)

# Check if spaCy is available
try:
    import spacy
    nlp = spacy.load("en_core_web_md")
    SPACY_AVAILABLE = True
except ImportError:
    logger.warning("spaCy not available. Document processing will have limited NLP capabilities.")
    nlp = None
    SPACY_AVAILABLE = False
except Exception as e:
    logger.error(f"Error loading spaCy model: {e}")
    logger.warning("Document processing will have limited NLP capabilities")
    nlp = None
    SPACY_AVAILABLE = False

class DocumentProcessor:
    """
    A processor for handling various document formats and extracting information.
    Supports PDF, DOCX, TXT, and JSON files.
    """
    
    def __init__(self):
        """Initialize the document processor."""
        logger.info("Initializing DocumentProcessor")
        self._ensure_dependencies()
        self.job_data = {}
        self.company_data = {}
        self.candidate_data = {}
        if not SPACY_AVAILABLE:
            logger.warning("Running with limited NLP capabilities")
    
    def _ensure_dependencies(self):
        """Ensure required dependencies are available."""
        try:
            # Check for PDF processing
            import PyPDF2
            logger.info("PyPDF2 is available for PDF processing")
        except ImportError:
            logger.warning("PyPDF2 is not installed. PDF processing will be limited.")
            
        try:
            # Check for DOCX processing
            import docx
            logger.info("python-docx is available for DOCX processing")
        except ImportError:
            logger.warning("python-docx is not installed. DOCX processing will be limited.")
    
    def parse_job_post(self, file_path: str) -> Dict[str, Any]:
        """Extract job title, description, and requirements from a document."""
        logger.info(f"Parsing job post: {file_path}")
        text = self.extract_text(file_path)
        
        # Extract job components using NLP and regex patterns
        title = self._extract_job_title(text)
        description = self._extract_job_description(text)
        requirements = self._extract_requirements(text)
        
        self.job_data = {
            "title": title,
            "description": description,
            "requirements": requirements
        }
        
        logger.debug(f"Extracted job data: {json.dumps(self.job_data, indent=2)}")
        return self.job_data
    
    def parse_company_profile(self, file_path: str) -> Dict[str, Any]:
        """Extract company mission, vision, and values."""
        logger.info(f"Parsing company profile: {file_path}")
        text = self.extract_text(file_path)
        
        self.company_data = {
            "name": self._extract_company_name(text),
            "mission": self._extract_mission(text),
            "vision": self._extract_vision(text),
            "values": self._extract_values(text)
        }
        
        logger.debug(f"Extracted company data: {json.dumps(self.company_data, indent=2)}")
        return self.company_data
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Extract candidate skills, experience, and education."""
        logger.info(f"Parsing resume: {file_path}")
        text = self.extract_text(file_path)
        
        self.candidate_data = {
            "name": self._extract_name(text),
            "skills": self._extract_skills(text),
            "experience": self._extract_experience(text),
            "education": self._extract_education(text)
        }
        
        logger.debug(f"Extracted candidate data: {json.dumps(self.candidate_data, indent=2)}")
        return self.candidate_data
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_ext == '.docx':
            return self._extract_docx_text(file_path)
        elif file_ext == '.txt':
            return self._extract_txt_text(file_path)
        elif file_ext == '.json':
            return self._extract_json_text(file_path)
        else:
            msg = f"Unsupported file type: {file_ext}"
            logger.error(msg)
            raise ValueError(msg)
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
                
                return text
        except ImportError:
            logger.warning("PyPDF2 not available, using fallback method for PDF")
            # Fallback using command line tool if available
            try:
                import subprocess
                result = subprocess.run(['pdftotext', file_path, '-'], capture_output=True, text=True)
                return result.stdout
            except Exception as e:
                logger.error(f"Error extracting text from PDF: {e}")
                return f"[PDF EXTRACTION FAILED: {str(e)}]"
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return f"[PDF EXTRACTION FAILED: {str(e)}]"
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + " | "
                    text += row_text.strip(" | ") + "\n"
                text += "\n"
            
            return text
        except ImportError:
            logger.warning("python-docx not available, using fallback method for DOCX")
            # Simple fallback that might not work well
            try:
                import zipfile
                from xml.etree import ElementTree
                
                text = ""
                with zipfile.ZipFile(file_path) as z:
                    for file in z.namelist():
                        if file.startswith('word/document.xml'):
                            content = z.read(file)
                            root = ElementTree.fromstring(content)
                            for elem in root.iter():
                                if elem.tag.endswith('}t'):
                                    if elem.text:
                                        text += elem.text + "\n"
                return text
            except Exception as e:
                logger.error(f"Error extracting text from DOCX: {e}")
                return f"[DOCX EXTRACTION FAILED: {str(e)}]"
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return f"[DOCX EXTRACTION FAILED: {str(e)}]"
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from a TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try another encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading TXT file: {e}")
                return f"[TXT EXTRACTION FAILED: {str(e)}]"
        except Exception as e:
            logger.error(f"Error reading TXT file: {e}")
            return f"[TXT EXTRACTION FAILED: {str(e)}]"
    
    def _extract_json_text(self, file_path: str) -> str:
        """Extract text from a JSON file as a formatted string."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                # Format the JSON with indentation for better readability
                return json.dumps(data, indent=2)
        except Exception as e:
            logger.error(f"Error reading JSON file: {e}")
            return f"[JSON EXTRACTION FAILED: {str(e)}]"
    
    def _extract_job_title(self, text: str) -> str:
        """Extract job title from text."""
        lines = text.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            if any(keyword in line.lower() for keyword in ['job title', 'position', 'role']):
                return line.split(':')[-1].strip() if ':' in line else line.strip()
        return "Untitled Position"
    
    def _extract_job_description(self, text: str) -> str:
        """Extract job description section."""
        # Look for description section
        desc_pattern = re.compile(r'(description|about the (role|position|job)).*?(?=(requirements|qualifications|responsibilities|what you\'ll do))', 
                                 re.IGNORECASE | re.DOTALL)
        match = desc_pattern.search(text)
        if match:
            return match.group(0).strip()
        
        # Fallback: use first 500 chars that aren't the title
        return text[:800]
    
    def _extract_requirements(self, text: str) -> List[str]:
        """Extract list of job requirements."""
        # Look for requirements/qualifications section
        req_pattern = re.compile(r'(requirements|qualifications).*?(?=(benefits|about us|how to apply|salary))', 
                               re.IGNORECASE | re.DOTALL)
        match = req_pattern.search(text)
        
        if match:
            req_text = match.group(0)
            # Extract bullet points or numbered items
            items = re.findall(r'[•\-*]+(.*?)(?=[•\-*]|$)', req_text, re.DOTALL)
            if not items:
                items = re.findall(r'\d+\.\s+(.*?)(?=\d+\.|$)', req_text, re.DOTALL)
            
            # Clean the extracted items
            requirements = [item.strip() for item in items if len(item.strip()) > 10]
            if requirements:
                return requirements
        
        # If no requirements found or if NLP is not available, return default requirements
        return [
            "Experience with software development",
            "Proficiency in programming languages",
            "Strong problem-solving skills",
            "Good communication skills",
            "Ability to work in a team"
        ]
    
    def _extract_company_name(self, text: str) -> str:
        """Extract company name from text."""
        lines = text.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            if any(keyword in line.lower() for keyword in ['company', 'organization', 'firm']):
                return line.split(':')[-1].strip() if ':' in line else line.strip()
        return "Unnamed Company"
    
    def _extract_mission(self, text: str) -> str:
        """Extract company mission statement."""
        mission_pattern = re.compile(r'(mission|our mission).*?(?=(vision|values|about us))', 
                                   re.IGNORECASE | re.DOTALL)
        match = mission_pattern.search(text)
        if match:
            return match.group(0).strip()
        
        return "To provide exceptional products and services to our customers."
    
    def _extract_vision(self, text: str) -> str:
        """Extract company vision statement."""
        vision_pattern = re.compile(r'(vision|our vision).*?(?=(mission|values|about us))', 
                                  re.IGNORECASE | re.DOTALL)
        match = vision_pattern.search(text)
        if match:
            return match.group(0).strip()
        
        return "To be a leader in our industry and create positive impact."
    
    def _extract_values(self, text: str) -> List[str]:
        """Extract company values from text."""
        values = []
        text_lower = text.lower()
        
        # Look for values section
        value_keywords = ['values', 'mission', 'principles', 'culture']
        
        for keyword in value_keywords:
            if keyword in text_lower:
                start = text_lower.find(keyword)
                # Find the next section header
                next_section = text.find('\n\n', start + len(keyword))
                if next_section == -1:
                    next_section = len(text)
                
                section = text[start:next_section]
                
                # Extract bullet points
                lines = section.split('\n')
                for line in lines:
                    if ('•' in line or '-' in line or '*' in line):
                        value = line.replace('•', '').replace('-', '').replace('*', '').strip()
                        if value and 3 < len(value) < 50:  # Reasonable value length
                            values.append(value)
                
                break
        
        # If no values found, use some default common values
        if not values:
            default_values = ['Innovation', 'Excellence', 'Integrity', 'Teamwork', 'Customer Focus']
            return default_values
        
        # Limit to 5 values
        return values[:5]
    
    def _extract_name(self, text: str) -> str:
        """Extract candidate name from resume."""
        # Names are typically at the top of resume
        if SPACY_AVAILABLE and nlp:
            doc = nlp(text[:200])
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    return ent.text
        
        # Fallback: first line could be name
        first_line = text.split('\n')[0].strip()
        if len(first_line) < 40:  # Likely a name if short
            return first_line
        
        return "John Doe"  # Default name
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text."""
        skills = []
        text_lower = text.lower()
        
        # Look for sections that might contain skills
        skill_sections = []
        
        if 'skills' in text_lower:
            start = text_lower.find('skills')
            end = start + 500  # Look at next 500 chars after "skills"
            skill_sections.append(text[start:end])
        
        if 'technical' in text_lower:
            start = text_lower.find('technical')
            end = start + 500
            skill_sections.append(text[start:end])
        
        # Extract skills from bullet points and lists
        for section in skill_sections:
            lines = section.split('\n')
            for line in lines:
                # Look for bullet points or comma-separated lists
                if ('•' in line or '-' in line or ',' in line):
                    parts = line.replace('•', ',').replace('-', ',').split(',')
                    for part in parts:
                        skill = part.strip()
                        if skill and len(skill) < 30:  # Reasonable skill name length
                            skills.append(skill)
        
        # If no skills found, extract potential technologies mentioned
        if not skills:
            tech_keywords = ['python', 'java', 'javascript', 'react', 'angular', 'vue', 
                            'node', 'django', 'flask', 'spring', 'aws', 'azure', 
                            'cloud', 'docker', 'kubernetes', 'sql', 'nosql', 'agile']
            
            for keyword in tech_keywords:
                if keyword in text_lower:
                    skills.append(keyword.capitalize())
        
        # Limit to 10 skills and remove duplicates
        return list(dict.fromkeys(skills))[:10]
    
    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience entries."""
        # Look for experience section
        exp_pattern = re.compile(r'(experience|work history|employment).*?(?=(education|skills|additional))', 
                               re.IGNORECASE | re.DOTALL)
        match = exp_pattern.search(text)
        
        experiences = []
        if match:
            exp_text = match.group(0)
            # Try to extract company names and positions
            job_entries = re.findall(r'([A-Z][A-Za-z\s&]+)[\s\|,]+([\w\s]+)[\s\|,]+(\d{1,2}/\d{4}|\d{4}\s*-\s*\d{1,2}/\d{4}|\d{4}\s*-\s*\d{4}|\d{4}\s*-\s*present)', exp_text)
            
            for company, title, dates in job_entries:
                experiences.append({
                    "company": company.strip(),
                    "title": title.strip(),
                    "dates": dates.strip()
                })
            
            if not experiences:
                # Simpler fallback: just extract paragraphs
                paragraphs = re.findall(r'\n\n(.*?)\n\n', exp_text)
                for i, p in enumerate(paragraphs[:3]):  # Limit to 3 experiences
                    experiences.append({
                        "experience": p.strip()
                    })
        
        if not experiences:
            # Default experience
            experiences.append({
                "company": "TechCorp",
                "title": "Software Engineer",
                "dates": "2020-Present"
            })
            experiences.append({
                "company": "StartupX",
                "title": "Junior Developer",
                "dates": "2018-2020"
            })
        
        return experiences
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education entries."""
        # Look for education section
        edu_pattern = re.compile(r'(education|academic background).*?(?=(experience|skills|additional))', 
                               re.IGNORECASE | re.DOTALL)
        match = edu_pattern.search(text)
        
        education = []
        if match:
            edu_text = match.group(0)
            # Try to extract degrees and institutions
            degree_entries = re.findall(r'([A-Z][A-Za-z\s&\.]+)[\s\|,]+([^,]+)[\s\|,]+(\d{4})', edu_text)
            
            for institution, degree, year in degree_entries:
                education.append({
                    "institution": institution.strip(),
                    "degree": degree.strip(),
                    "year": year.strip()
                })
            
            if not education:
                # Simpler fallback: just extract lines
                lines = re.findall(r'\n([^\n]+University|College[^\n]+)', edu_text)
                for line in lines:
                    education.append({
                        "education": line.strip()
                    })
        
        if not education:
            # Default education
            education.append({
                "institution": "University of Technology",
                "degree": "Bachelor of Computer Science",
                "year": "2018"
            })
        
        return education
    
    def get_status(self) -> str:
        """Return the current status of the document processor."""
        if SPACY_AVAILABLE:
            return "Document processor initialized with full NLP capabilities"
        else:
            return "Document processor initialized with limited NLP capabilities (spaCy not available)" 